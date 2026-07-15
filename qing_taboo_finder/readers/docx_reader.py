"""DOCX reader."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from ..models import LoadedDocument


def read_docx(path: Path) -> LoadedDocument:
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(part for part in paragraphs if part)
    return LoadedDocument(path=path, text=text, title=path.stem)